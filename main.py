import os
import re
import time
from google import generativeai as genai
from google.generativeai.types import HarmCategory, HarmBlockThreshold
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.section import WD_ORIENT
from docx.enum.style import WD_STYLE_TYPE

def sanitize_filename(text):
    """Sanitize the filename by removing invalid characters."""
    return re.sub(r'[\\/*?:"<>|]', "", text.strip()).replace(" ", "_")

def add_code_style(document):
    """Add a custom style for code formatting."""
    styles = document.styles
    style = styles.add_style('Code', WD_STYLE_TYPE.PARAGRAPH)
    style.font.name = 'Courier New'
    style.font.size = Pt(10)
    style.paragraph_format.space_after = Pt(0)
    style.paragraph_format.space_before = Pt(0)


def style_code(document):
    """Apply the 'Code' style to all paragraphs in the document."""
    for paragraph in document.paragraphs:
        if paragraph.style.name == 'Code':
            # Ensure pPr exists
            pPr = paragraph._element.get_or_add_pPr()
            # Set background color to black
            shading = OxmlElement('w:shd')
            shading.set(qn('w:fill'), '000000')  # Black background
            pPr.append(shading)

            # Set font color to green
            for run in paragraph.runs:
                run.font.color.rgb = RGBColor(0, 255, 0)  # Green font

# ========== TOPIC LIST ==========
def create_topic_list(text, main_context):
    gemini_api_key = "AIzaSyA1S7mQPuZQR6oERynreZhhD0bd8Fizzzc"
    genai.configure(api_key=gemini_api_key)

    generation_config = {
        'temperature': 0.1,
        'stop_sequences': [],
    }

    safety_settings = [
        {"category": "HARM_CATEGORY_HARASSMENT", "threshold": "BLOCK_ONLY_HIGH"},
        {"category": "HARM_CATEGORY_HATE_SPEECH", "threshold": "BLOCK_ONLY_HIGH"},
        {"category": "HARM_CATEGORY_SEXUALLY_EXPLICIT", "threshold": "BLOCK_ONLY_HIGH"},
        {"category": "HARM_CATEGORY_DANGEROUS_CONTENT", "threshold": "BLOCK_ONLY_HIGH"}
    ]

    model = genai.GenerativeModel('gemini-2.0-flash', generation_config=generation_config, safety_settings=safety_settings)

    query = f"""
    Gere uma lista de tópicos sobre o conteúdo: "{text}", dentro do contexto global: "{main_context}".
    A lista deve seguir o padrão:
    1- Nome do Tópico 1
    2- Nome do Tópico 2
    3- Nome do Tópico 3
    """

    response = model.generate_content(query)
    output = response.text

    with open("topics.txt", "w", encoding="utf-8") as text_file:
        text_file.write(output)

# ========== GERAÇÃO DE TEXTO ==========
def generate_topic(main_context, text):
    gemini_api_key = "AIzaSyA1S7mQPuZQR6oERynreZhhD0bd8Fizzzc"
    genai.configure(api_key=gemini_api_key)

    generation_config = {
        'temperature': 0.9,
        'top_p': 1,
        'top_k': 40,
        'stop_sequences': [],
    }

    safety_settings = [
        {"category": "HARM_CATEGORY_HARASSMENT", "threshold": "BLOCK_ONLY_HIGH"},
        {"category": "HARM_CATEGORY_HATE_SPEECH", "threshold": "BLOCK_ONLY_HIGH"},
        {"category": "HARM_CATEGORY_SEXUALLY_EXPLICIT", "threshold": "BLOCK_ONLY_HIGH"},
        {"category": "HARM_CATEGORY_DANGEROUS_CONTENT", "threshold": "BLOCK_ONLY_HIGH"}
    ]

    model = genai.GenerativeModel('gemini-2.0-flash', generation_config=generation_config, safety_settings=safety_settings)

    query = f"""
    Gere o texto para um documento técnico sobre o seguinte tópico: "{text}". 
    O contexto global do documento é: "{main_context}".

    O texto deve estar em português de Portugal, ser claro, objetivo e completo, com no mínimo 600 palavras e no máximo 1200 palavras.

    Estruture o conteúdo com as seguintes marcações:
    - Use `##` para o título principal.
    - Use `###` para subtítulos.
    - Separe os parágrafos com linhas em branco.
    """

    response = model.generate_content(query)
    time.sleep(20)  # Aumentando o tempo de espera para evitar exceder a cota
    return response.text

# ========== EXECUTA TUDO ==========
def perform_creation(main_context):
    file_path = "topics.txt"
    output_dir = "textos"
    os.makedirs(output_dir, exist_ok=True)

    collected_texts = []

    with open(file_path, "r", encoding="utf-8") as text_file:
        for line in text_file:
            topic = line.strip()
            if not topic:
                continue
            print(f"🔹 Gerando texto: {topic}")
            try:
                content = generate_topic(main_context, topic)
                filename = sanitize_filename(topic) + ".txt"
                filepath = os.path.join(output_dir, filename)

                with open(filepath, "w", encoding="utf-8") as f:
                    f.write(content)

                collected_texts.append((topic, content))
            except Exception as e:
                print(f"Erro ao gerar '{topic}': {e}")
                continue

    export_docx(collected_texts)
    print("🎉 Finalizado com sucesso!")
def read_collected_texts(output_dir="textos"):
    collected_texts = []
    files = sorted(os.listdir(output_dir), key=lambda x: int(re.search(r'\d+', x).group()) if re.search(r'\d+', x) else float('inf'))
    for filename in files:
        if filename.endswith(".txt"):
            topic = filename[:-4]  # Remove .txt extension
            with open(os.path.join(output_dir, filename), "r", encoding="utf-8") as f:
                content = f.read()
                collected_texts.append((topic, content))
    return collected_texts
# ========== CRIA DOCX CHIC ==========
def export_docx(pages, output_path="data/trabalho_f6_som.docx"):
    document = Document()
    add_code_style(document)  # Adiciona o estilo 'Code' ao documento

    # CAPA
    section = document.sections[0]
    section.page_height = Inches(11.69)
    section.page_width = Inches(8.27)
    section.orientation = WD_ORIENT.PORTRAIT
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)

    # Adiciona o título do documento
    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Trabalho de Grupo F6 - Som")
    run.bold = True
    run.font.size = Pt(28)

    # Adiciona o subtítulo do documento
    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Caleb Gomes Pinto, Danylo Petrash")
    run.italic = True
    run.font.size = Pt(18)

    document.add_paragraph().add_run().add_break()
    document.add_page_break()

    # SUMÁRIO
    toc_title = document.add_paragraph("Sumário")
    toc_title.style = 'Heading 1'
    fldSimple = OxmlElement('w:fldSimple')
    fldSimple.set(qn('w:instr'), 'TOC \\o "1-3" \\h \\z \\u')
    document._body._element.append(fldSimple)
    document.add_page_break()

    # CONTEÚDO
    for index, (title, content) in enumerate(pages, start=1):
        lines = content.strip().split("\n")
        in_code_block = False
        code_lines = []

        for line in lines:
            line = line.strip()
            if not line:
                continue
            elif line.startswith("## "):
                document.add_heading(line.replace("## ", "").strip(), level=1)
            elif line.startswith("### "):
                document.add_heading(line.replace("### ", "").strip(), level=2)
            elif line.startswith("#### "):
                document.add_heading(line.replace("#### ", "").strip(), level=3)
            elif line.startswith("##### "):
                document.add_heading(line.replace("##### ", "").strip(), level=4)
            elif line.startswith("```") and not in_code_block:
                in_code_block = True
                code_lines = []  # Reinicia a coleta de linhas de código
            elif line.startswith("```") and in_code_block:
                in_code_block = False
                # Adiciona as linhas de código coletadas como um único parágrafo
                document.add_paragraph("\n".join(code_lines), style='Code')
            elif in_code_block:
                code_lines.append(line)  # Coleta linhas para o bloco de código
            else:
                # Processa o texto para aplicar formatação em negrito e listas
                process_bold_and_list_text(document, line)


    # CABEÇALHO
    for section in document.sections:
        header = section.header.paragraphs[0]
        header.text = "Trabalho de Grupo F6 - Som"
        header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    style_code(document)  # Aplica estilo de código

    # SALVAR
    os.makedirs("data", exist_ok=True)
    document.save(output_path)
    print(f"📄 Documento Word salvo em: {output_path}")

def process_bold_and_list_text(document, line):
    """
    Processa uma linha de texto para aplicar formatação em negrito
    onde o texto está entre asteriscos duplos (**), e para criar
    itens de lista onde o texto começa com um asterisco simples (*).
    
    Parâmetros:
    document (Document): O objeto do documento onde o texto será adicionado.
    line (str): A linha de texto a ser processada.
    """
    # Verifica se a linha começa com um asterisco simples para lista
    if line.startswith("* "):
        # Remove o asterisco e o espaço para o texto da lista
        line = line[2:]  # Remove "* "
        list_item = document.add_paragraph(style='ListBullet')  # Adiciona um item de lista
    else:
        list_item = document.add_paragraph()  # Adiciona um parágrafo normal

    parts = line.split("**")
    for i, part in enumerate(parts):
        if i % 2 == 1:  # Se o índice é ímpar, aplica negrito
            run = list_item.add_run(part)
            run.bold = True
        else:
            list_item.add_run(part)  # Adiciona texto normal

if __name__ == "__main__":
    export_docx(read_collected_texts())